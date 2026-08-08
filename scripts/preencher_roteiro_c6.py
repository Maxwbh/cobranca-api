#!/usr/bin/env python3
"""Escreve a evidência real do sandbox no formulário de homologação do C6.

O `.docx` do banco tem 68 tabelas vazias com "Status Code - Retornado" e
"Response Body - Retornado" — uma por etapa de caso. Preencher à mão é copiar
JSON 68 vezes, e é o que mantém a homologação parada.

Pipeline:

    PYTHONPATH=gateway python scripts/homologacao_c6.py --json > evidencia.json
    python scripts/preencher_roteiro_c6.py evidencia.json \\
        --roteiro "Roteiro de Testes - C6 Developers v3.0.docx"

Sai `<roteiro> - preenchido.docx`, com **só o que a execução devolveu**. Caso
sem resultado no JSON fica em branco — este script não inventa retorno, e é
por isso que ele lê de um arquivo de evidência em vez de aceitar valores na
linha de comando.

Além das tabelas de retorno, o script fecha as duas áreas de formulário que
o banco confere primeiro e que não saem da evidência:

- **DADOS DA ORGANIZAÇÃO** — os seis campos vêm de `--cnpj`, `--empresa`,
  `--software`, `--responsavel`, `--email` e `--telefone`. São dados
  cadastrais, não resultado de teste, então entram por argumento mesmo.
- **MARQUE O CHECKBOX COM AS APIS QUE VOCÊ TESTARÁ** — marcado a partir da
  própria evidência: API com ao menos um caso executado é marcada; API cujos
  casos são todos "ausente" (fora de escopo) fica desmarcada, porque o banco
  manda **não** marcar o que a aplicação não usa.

Baixe o formulário em:
https://developers.c6bank.com.br/test-scripts/Roteiro%20de%20Testes%20-%20C6%20Developers%20v3.0.docx
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    import docx
except ImportError:
    sys.exit("falta a dependência: pip install python-docx")

CASO_RE = re.compile(r"([A-Z]{1,3})_(\d{2})_?(\d{2})?")
LIMITE_CORPO = 4000  # corpo gigante estoura a tabela do Word
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Nome da API na tabela de checkboxes → prefixo dos casos na evidência.
# Casado por igualdade exata do texto da célula: "PIX" é substring de
# "PIX AUTOMÁTICO", e um `in` marcaria a caixa errada.
API_PREFIXO = {
    "AGENDAMENTO DE PAGAMENTOS": "AP",
    "BOLETO": "B",
    "CHECKOUT": "C",
    "EXTRATO": "E",
    "PIX": "P",
    "TRANSAÇÕES E RECEBÍVEIS": "TR",
    "PIX AUTOMÁTICO": "PA",
    "BOLEPIX": "BP",
}

# O formulário v3.0 do C6 rotula duas tabelas com o identificador da seção
# anterior. O título da seção acima delas não deixa dúvida sobre o que são, e
# casar pelo rótulo faria a segunda tabela disputar a evidência da primeira —
# as duas erradas. Corrigido por índice, que é o que identifica a tabela de
# fato. Se o banco corrigir o formulário, o aviso de repetidos some e este
# mapa deixa de casar (o índice muda) — falha visível, não silenciosa.
ROTULO_ERRADO = {
    75: "P_03_03",  # seção "P_03 – COBRANÇA PIX EM LOTE": consultar lote específico
    85: "P_04_04",  # seção "P_04 – GERENCIAMENTO DE LOCATION": desvincular cobrança
}


def tabelas_de_retorno(doc) -> list[tuple[int, str, int, object]]:
    """(índice, caso, linha_a_preencher, tabela) de cada tabela 'Retornado'.

    O formulário usa DUAS formas: a maioria tem 3 linhas (título / cabeçalho /
    células), mas Pix Automático e BP_01 têm 4 — cabe um subtítulo entre o
    título e o cabeçalho. Fixar o cabeçalho na linha 1, como estava, tornava
    essas 17 tabelas invisíveis: elas não eram preenchidas e nem entravam na
    contagem, então o relatório dizia "51 de 51" com 17 em branco fora da conta.
    Silêncio é o pior modo de falhar num documento que atesta teste.

    Procura o cabeçalho em qualquer linha e preenche a seguinte.
    """
    achados = []
    for i, t in enumerate(doc.tables):
        if len(t.columns) < 2:
            continue
        cabecalho = next((r for r, linha in enumerate(t.rows)
                          if "Retornado" in linha.cells[0].text), None)
        if cabecalho is None or cabecalho + 1 >= len(t.rows):
            continue
        m = CASO_RE.search(t.rows[0].cells[0].text)
        # "C_0501" é C_05_01 sem o separador — o bloco de checkout escreve
        # assim. Sem normalizar, as duas tabelas viram "C_05" e colidem.
        caso = ""
        if m:
            caso = f"{m.group(1)}_{m.group(2)}" + (f"_{m.group(3)}" if m.group(3) else "")
        achados.append((i, ROTULO_ERRADO.get(i, caso), cabecalho + 1, t))
    return achados


def escrever(celula, texto: str) -> None:
    """Substitui o conteúdo preservando a formatação do primeiro run."""
    p = celula.paragraphs[0]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(texto)
    for extra in celula.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def preencher_identificacao(doc, campos: dict[str, str]) -> list[str]:
    """Preenche a tabela DADOS DA ORGANIZAÇÃO; devolve os rótulos preenchidos.

    Casa cada linha pelo rótulo da segunda coluna, não pela posição: se o
    banco reordenar os campos numa próxima versão, o valor continua indo para
    a linha certa — ou não vai, e o relatório final acusa o campo que sobrou.
    """
    tabela = next((t for t in doc.tables
                   if t.rows and "DADOS DA ORGANIZAÇÃO" in t.rows[0].cells[0].text), None)
    if tabela is None:
        print("aviso: tabela DADOS DA ORGANIZAÇÃO não encontrada — identificação não preenchida")
        return []
    rotulo_valor = {
        "CNPJ": campos.get("cnpj"),
        "NOME DA EMPRESA": campos.get("empresa"),
        "NOME DO SOFTWARE": campos.get("software"),
        "RESPONSÁVEL": campos.get("responsavel"),
        "E-MAIL": campos.get("email"),
        "TELEFONE": campos.get("telefone"),
    }
    feitos = []
    for linha in tabela.rows[1:]:
        rotulo = linha.cells[1].text.strip()
        valor = next((v for chave, v in rotulo_valor.items() if chave in rotulo), None)
        if valor:
            escrever(linha.cells[-1], valor)
            feitos.append(rotulo)
    return feitos


def marcar_checkboxes(doc, exercitadas: set[str]) -> tuple[list[str], list[str]]:
    """Marca o checkbox de cada API com evidência; devolve (marcadas, desmarcadas).

    Os checkboxes são campos legados (FORMCHECKBOX) e o estado mora no
    `w:checked` dentro do `w:ffData` — python-docx não tem API para isso,
    então é XML mesmo. Os nomes dos campos não servem de chave: o formulário
    v3.0 repete "Check6" três vezes, então o casamento é pela linha da tabela
    em que o checkbox está, via API_PREFIXO.
    """
    tabela = next((t for t in doc.tables
                   if t.rows and "MARQUE O CHECKBOX" in t.rows[0].cells[0].text), None)
    if tabela is None:
        print("aviso: tabela de checkboxes não encontrada — nenhuma API marcada")
        return [], []
    marcadas, desmarcadas = [], []
    for linha in tabela.rows:
        textos = {c.text.strip() for c in linha.cells}
        api = next((nome for nome in API_PREFIXO if nome in textos), None)
        caixas = linha._tr.findall(f".//{W}checkBox")
        if api is None or not caixas:
            continue
        marcar = API_PREFIXO[api] in exercitadas
        for caixa in caixas:
            checked = caixa.find(f"{W}checked")
            if checked is None:
                checked = caixa.makeelement(f"{W}checked", {})
                caixa.append(checked)
            checked.set(f"{W}val", "1" if marcar else "0")
        (marcadas if marcar else desmarcadas).append(api)
    return marcadas, desmarcadas


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("evidencia", help="JSON gerado por homologacao_c6.py --json")
    ap.add_argument("--roteiro", required=True, help="o .docx baixado do portal do C6")
    ap.add_argument("--saida", help="default: '<roteiro> - preenchido.docx'")
    ident = ap.add_argument_group("identificação (tabela DADOS DA ORGANIZAÇÃO)")
    ident.add_argument("--cnpj")
    ident.add_argument("--empresa")
    ident.add_argument("--software")
    ident.add_argument("--responsavel")
    ident.add_argument("--email")
    ident.add_argument("--telefone")
    a = ap.parse_args()

    dados = json.loads(Path(a.evidencia).read_text(encoding="utf-8"))
    por_caso: dict[str, list[dict]] = {}
    for r in dados.get("resultados", []):
        por_caso.setdefault(r["caso"], []).append(r)
    # API exercitada = ao menos um caso com resultado real; prefixo cujos
    # casos são todos "ausente" está fora do produto e fica desmarcado.
    exercitadas = {caso.split("_")[0] for caso, fila in por_caso.items()
                   if any(not r.get("ausente") for r in fila)}

    doc = docx.Document(a.roteiro)
    alvos = tabelas_de_retorno(doc)

    campos = {k: getattr(a, k) for k in
              ("cnpj", "empresa", "software", "responsavel", "email", "telefone")}
    identificados = preencher_identificacao(doc, campos) if any(campos.values()) else []
    marcadas, desmarcadas = marcar_checkboxes(doc, exercitadas)

    # ROTULO_ERRADO já corrige as duas repetições conhecidas do formulário v3.0.
    # O aviso continua para o caso de o banco introduzir outras — casar só pelo
    # ID faria a segunda tabela disputar a evidência da primeira e ficar em
    # branco sem explicação.
    vistos: dict[str, int] = {}
    for _, caso, _lin, _t in alvos:
        vistos[caso] = vistos.get(caso, 0) + 1
    repetidos = {c: n for c, n in vistos.items() if n > 1 and c}
    if repetidos:
        print("aviso: identificadores repetidos no formulário do banco — "
              + ", ".join(f"{c}×{n}" for c, n in sorted(repetidos.items())))
        print("       a evidência vai para a primeira ocorrência; confira as demais à mão")

    usados: dict[str, int] = {}
    preenchidos, vazios, ausentes = [], [], []
    for _, caso, linha, tabela in alvos:
        fila = por_caso.get(caso, [])
        n = usados.get(caso, 0)
        if n >= len(fila):          # etapa sem evidência: fica em branco, de propósito
            vazios.append(caso or "?")
            continue
        usados[caso] = n + 1
        r = fila[n]
        if r.get("ausente"):
            # Caso fora do produto não é falha nem lacuna: é fronteira. Escrever
            # o motivo é mais honesto que deixar a célula vazia, que se lê como
            # "esqueceram" — e o banco pede justamente que o não-usado se declare.
            escrever(tabela.rows[linha].cells[0], "N/A")
            escrever(tabela.rows[linha].cells[1], f"Não se aplica — {r.get('motivo', '')}")
            ausentes.append(caso)
            continue
        status = "" if r.get("status_code") is None else str(r["status_code"])
        if not status:
            status = "2xx" if r.get("ok") else "erro"
        corpo = r.get("response_body")
        texto = corpo if isinstance(corpo, str) else json.dumps(
            corpo, ensure_ascii=False, indent=2, default=str)
        if len(texto) > LIMITE_CORPO:
            texto = (texto[:LIMITE_CORPO]
                     + f"\n[TRUNCADO — corpo completo tem {len(texto)} caracteres; "
                       f"exibidos os {LIMITE_CORPO} primeiros]")
        if not r.get("ok"):
            texto = f"{r.get('erro','')}\n\n{texto}"
        escrever(tabela.rows[linha].cells[0], status)
        escrever(tabela.rows[linha].cells[1], texto)
        preenchidos.append(caso)

    saida = a.saida or str(Path(a.roteiro).with_suffix("")) + " - preenchido.docx"
    doc.save(saida)

    print(f"gerado: {saida}")
    if identificados:
        print(f"  identificação: {len(identificados)} de 6 campos preenchidos")
        faltou = 6 - len(identificados)
        if faltou:
            print(f"  aviso: {faltou} campo(s) da identificação sem valor ou sem linha correspondente")
    elif any(campos.values()):
        print("  aviso: identificação pedida mas nenhum campo casou com a tabela")
    else:
        print("  identificação: não preenchida (passe --cnpj, --empresa, … para preencher)")
    if marcadas or desmarcadas:
        print(f"  checkboxes : marcadas {', '.join(marcadas)}"
              + (f" | desmarcadas {', '.join(desmarcadas)}" if desmarcadas else ""))
    print(f"  preenchidas: {len(preenchidos)} de {len(alvos)} tabelas")
    if ausentes:
        print(f"  ausentes   : {len(ausentes)} — marcadas N/A com o motivo "
              f"({', '.join(sorted(set(ausentes)))})")
    if vazios:
        faltam = sorted(set(vazios))
        print(f"  em branco  : {len(vazios)} — {', '.join(faltam)}")
        print("  (etapa sem evidência no JSON fica em branco; rode o caso e refaça)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
